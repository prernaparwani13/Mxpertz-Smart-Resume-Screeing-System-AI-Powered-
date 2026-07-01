import os
from docx import Document

os.makedirs('samples', exist_ok=True)

# 1. Sample Job Description
jd_text = """
We are looking for a Backend Software Engineer.
Requirements:
- Strong experience with Python and FastAPI.
- Familiarity with SQL and PostgreSQL databases.
- Experience with Git, Docker, and AWS.
- Knowledge of Machine Learning concepts is a plus.
- Agile and Scrum experience.
"""
with open('samples/sample_jd.txt', 'w') as f:
    f.write(jd_text.strip())

# 2. Sample Resume 1 (Good Match)
doc1 = Document()
doc1.add_heading('John Doe - Backend Developer', 0)
doc1.add_paragraph('Experienced backend developer with 5 years of experience in building scalable REST APIs.')
doc1.add_heading('Skills', level=1)
doc1.add_paragraph('Python, FastAPI, SQL, PostgreSQL, Git, Docker, Kubernetes.')
doc1.add_heading('Experience', level=1)
doc1.add_paragraph('Developed high-performance microservices using FastAPI and deployed them on AWS.')
doc1.save('samples/john_doe_backend.docx')

# 3. Sample Resume 2 (Poor Match)
doc2 = Document()
doc2.add_heading('Jane Smith - Frontend Developer', 0)
doc2.add_paragraph('Creative frontend developer specializing in modern user interfaces.')
doc2.add_heading('Skills', level=1)
doc2.add_paragraph('JavaScript, TypeScript, React, HTML, CSS, Tailwind, Bootstrap, Figma.')
doc2.add_heading('Experience', level=1)
doc2.add_paragraph('Built responsive web applications using React and Tailwind CSS. Integrated with various REST APIs.')
doc2.save('samples/jane_smith_frontend.docx')

print("Sample files created in the 'samples' directory.")
